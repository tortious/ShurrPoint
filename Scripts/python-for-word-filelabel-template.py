from docx import Document

# Load the template
doc = Document("path_to/Case_Template.docx")

# Replace placeholders with actual data
case_data = {
    "PLAINTIFF": "John Doe",
    "DEFENDANT": "Jane Smith",
    "case #": "12345",
    "claim #": "67890",
    "INSURANCE COMPANY": "FC Insurance",
    "ADJUSTER": "Mark",
    "ADJUSTER EMAIL": "mark@example.com",
    "CLIENT NAME": "Doe Co.",
    "CLIENT ADDRESS": "123 Main St",
    "PHONE": "555-1234",
    "EMAIL": "client@example.com",
    "OPPOSING PARTY": "Smith Inc.",
    "OPPOSING COUNSEL": "William Black",
    "OC ADDRESS": "456 Elm St",
    "OC PHONE": "555-5678",
    "OC EMAIL": "opposing@example.com"
}

# Function to replace text in document
def replace_text(doc, data):
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(f"[{key}]", value)

    # Update text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(f"[{key}]", value)

replace_text(doc, case_data)

# Save the filled document
doc.save("Filled_Case_Document.docx")
